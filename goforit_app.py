import streamlit as st
import plotly.graph_objects as go
import pandas as pd
import requests
import os
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches

# --- CONFIGURACIÓN DE PÁGINA ---
st.set_page_config(page_title="Go For It | Strategic Advisor", page_icon="🚀", layout="wide")

# --- ESTILO PERSONALIZADO ---
st.markdown("""
    <style>
    .main { background-color: #f8fafc; }
    [data-testid="stSidebar"] { background-color: #0f172a; color: white; }
    .stMetric { background-color: white; padding: 20px; border-radius: 12px; box-shadow: 0 4px 6px -1px rgb(0 0 0 / 0.1); }
    div.stButton > button { width: 100%; border-radius: 8px; font-weight: bold; background-color: #2563eb; color: white; }
    </style>
    """, unsafe_allow_html=True)

# --- CONFIGURACIÓN DE CONEXIÓN ---
SHEET_ID = st.secrets["SHEET_ID"]
SCRIPT_URL = st.secrets["SCRIPT_URL"]
LOGO_PATH = "Logos_GoForIt"
if not os.path.exists(LOGO_PATH): os.makedirs(LOGO_PATH, exist_ok=True)

# --- FUNCIONES CORE ---
@st.cache_data(ttl=5)
def leer_datos():
    try:
        url = f"https://docs.google.com/spreadsheets/d/{SHEET_ID}/export?format=csv"
        df = pd.read_csv(url).dropna(subset=['Empresa'])
        
        # NORMALIZACIÓN: Unificamos 'Distintos' y 'Capacidades_Distintivas' para que siempre cargue
        if "Distintos" in df.columns and "Capacidades_Distintivas" not in df.columns:
            df = df.rename(columns={"Distintos": "Capacidades_Distintivas"})
        elif "Capacidades_Distintivas" in df.columns and "Distintos" not in df.columns:
            df["Distintos"] = df["Capacidades_Distintivas"]
            
        cols_check = ["Capacidades_Distintivas", "Distintos", "Facilita", "Dificulta", "No_Hacemos", "Accion_1", "Accion_2"]
        for c in cols_check:
            if c not in df.columns: df[c] = ""
        return df
    except:
        return pd.DataFrame()

def enviar_datos(payload):
    try:
        res = requests.post(SCRIPT_URL, json=payload, timeout=15)
        return res.status_code == 200
    except: return False

def generar_word(empresa, df_f, logo_p):
    doc = Document()
    doc.add_heading(f'Informe Estratégico: {empresa}', 0)
    if os.path.exists(logo_p):
        try: doc.add_picture(logo_p, width=Inches(1.5))
        except: pass
    doc.add_heading('Detalle de Hallazgos y Capacidades', level=1)
    for _, r in df_f.iterrows():
        doc.add_heading(f"{r['Dimensión']} - {r['Foco']} ({r['Nombre']})", level=2)
        doc.add_paragraph(f"Madurez: {r['Actual']}/10 | Meta: {r['Objetivo']}/10")
        cap = r['Capacidades_Distintivas'] if pd.notna(r['Capacidades_Distintivas']) else r['Distintos']
        doc.add_paragraph(f"Capacidades Distintivas: {cap}").italic = True
        doc.add_paragraph(f"Acciones: 1. {r['Accion_1']} | 2. {r['Accion_2']}")
    target = BytesIO()
    doc.save(target)
    return target.getvalue()

# --- CARGA DE DATOS ---
df_global = leer_datos()
empresas_list = sorted(df_global['Empresa'].unique().tolist()) if not df_global.empty else []

# --- SIDEBAR: PROYECTOS ---
with st.sidebar:
    st.title("Consultoría Pro")
    emp_sel = st.selectbox("🎯 Seleccionar Proyecto:", ["-- Nuevo Proyecto --"] + empresas_list)
    
    if emp_sel == "-- Nuevo Proyecto --":
        st.subheader("✨ Registrar Cliente")
        n_emp = st.text_input("Nombre de la Empresa:")
        u_logo = st.file_uploader("Logo (Opcional):", type=["png", "jpg"])
        if st.button("🚀 Activar Proyecto"):
            if n_emp:
                if enviar_datos({"Empresa": n_emp, "Dimensión": "Alineación Estratégica", "Foco": "General", "Nombre": "Admin", "Actual": 1, "Objetivo": 1}):
                    if u_logo: Image.open(u_logo).save(os.path.join(LOGO_PATH, f"{n_emp}.png"))
                    st.cache_data.clear()
                    st.rerun()
        st.stop()

# --- DATOS DEL CLIENTE SELECCIONADO ---
df_emp = df_global[df_global['Empresa'] == emp_sel]
logo_file = os.path.join(LOGO_PATH, f"{emp_sel}.png")

# --- HEADER APP ---
c_h1, c_h2 = st.columns([1, 5])
with c_h1: 
    if os.path.exists(logo_file): st.image(logo_file, width=110)
with c_h2: st.title(f"Intervención: {emp_sel}")

tab1, tab2, tab3 = st.tabs(["📝 Registro y Edición", "📊 Radar Analítico", "📋 Reporte Ejecutivo"])

# --- TAB 1: REGISTRO Y CARGA DE DATOS ---
with tab1:
    st.markdown("### Configuración de la Evaluación")
    c1, c2, c3 = st.columns(3)
    dim_f = c1.selectbox("Pilar:", ["Intimidad Cliente-Mercado", "Liderazgo Producto-Servicio", "Excelencia Operacional", "Flujo de Caja Sostenible", "Aprendizaje Constante", "Alineación Estratégica"])
    
    m_foc = c2.toggle("Nuevo Foco")
    foc_f = c2.text_input("Nombre Foco:") if m_foc else c2.selectbox("Foco Actual:", sorted(df_emp['Foco'].unique()) if not df_emp.empty else ["General"])
    
    m_nom = c3.toggle("Nuevo Participante")
    nom_f = c3.text_input("Nombre Participante:") if m_nom else c3.selectbox("Participante:", sorted(df_emp['Nombre'].unique()) if not df_emp.empty else ["Admin"])

    # Búsqueda de datos previos para autocarga
    match = df_emp[(df_emp['Dimensión'] == dim_f) & (df_emp['Foco'] == foc_f) & (df_emp['Nombre'] == nom_f)]
    row = match.iloc[0] if not match.empty else None

    with st.form("form_smart_advisor"):
        st.subheader("✏️ Editar Datos" if row is not None else "➕ Crear Registro")
        
        # Sliders
        f_a, f_b = st.columns(2)
        v_act = f_a.slider("Nivel Actual (Hoy)", 1, 10, int(row['Actual']) if row is not None else 5)
        v_obj = f_b.slider("Nivel Meta (Objetivo)", 1, 10, int(row['Objetivo']) if row is not None else 8)
        
        # Función para extraer texto de forma segura (buscando en ambos nombres de columna)
        def get_text(r, col_p, col_alt):
            if r is None: return ""
            if col_p in r and pd.notna(r[col_p]): return str(r[col_p])
            if col_alt in r and pd.notna(r[col_alt]): return str(r[col_alt])
            return ""

        st.markdown("---")
        tx1, tx2 = st.columns(2)
        f_fac = tx1.text_area("¿Qué facilita el proceso?", value=get_text(row, 'Facilita', 'Facilita'))
        f_dif = tx1.text_area("¿Qué dificulta el proceso?", value=get_text(row, 'Dificulta', 'Dificulta'))
        f_no = tx2.text_area("¿Qué NO estamos haciendo?", value=get_text(row, 'No_Hacemos', 'No_Hacemos'))
        
        # Campo Crítico: Capacidades Distintivas
        f_cap = tx2.text_area("Capacidades Distintivas:", value=get_text(row, 'Capacidades_Distintivas', 'Distintos'), help="Habilidades o recursos únicos de la organización.")
        
        a1 = st.text_input("Acción Inmediata 1:", value=get_text(row, 'Accion_1', 'Accion_1'))
        a2 = st.text_input("Acción Inmediata 2:", value=get_text(row, 'Accion_2', 'Accion_2'))

        if st.form_submit_button("💾 Guardar y Sincronizar"):
            payload = {
                "Empresa": emp_sel, "Dimensión": dim_f, "Foco": foc_f, "Nombre": nom_f,
                "Actual": v_act, "Objetivo": v_obj, "Brecha": v_obj - v_act,
                "Facilita": f_fac, "Dificulta": f_dif, "No_Hacemos": f_no,
                "Capacidades_Distintivas": f_cap, "Distintos": f_cap, # Enviamos ambos por seguridad
                "Accion_1": a1, "Accion_2": a2
            }
            if enviar_datos(payload):
                st.success("¡Información actualizada en la nube!")
                st.cache_data.clear()
                st.rerun()

# --- FUNCIÓN DE FILTROS ---
def interfaz_filtros(df, k):
    st.markdown("### 🔍 Segmentar Análisis")
    cf1, cf2, cf3 = st.columns(3)
    p = cf1.multiselect("Pilares:", df['Dimensión'].unique(), default=df['Dimensión'].unique(), key=f"p_{k}")
    f = cf2.multiselect("Focos:", df['Foco'].unique(), default=df['Foco'].unique(), key=f"f_{k}")
    n = cf3.multiselect("Participantes:", df['Nombre'].unique(), default=df['Nombre'].unique(), key=f"n_{k}")
    return df[(df['Dimensión'].isin(p)) & (df['Foco'].isin(f)) & (df['Nombre'].isin(n))]

# --- TAB 2: RADAR ---
with tab2:
    if not df_emp.empty:
        df_f_rad = interfaz_filtros(df_emp, "radar")
        if not df_f_rad.empty:
            res = []
            for d in ["Intimidad Cliente-Mercado", "Liderazgo Producto-Servicio", "Excelencia Operacional", "Flujo de Caja Sostenible", "Aprendizaje Constante", "Alineación Estratégica"]:
                sub = df_f_rad[df_f_rad['Dimensión'] == d]
                res.append({"Dimensión": d, "Actual": sub['Actual'].mean() if not sub.empty else 0, "Objetivo": sub['Objetivo'].mean() if not sub.empty else 0})
            df_p = pd.DataFrame(res); df_p["Brecha"] = df_p["Objetivo"] - df_p["Actual"]
            
            fig = go.Figure()
            fig.add_trace(go.Scatterpolar(r=df_p['Objetivo'], theta=df_p['Dimensión'], fill='toself', name='Meta'))
            fig.add_trace(go.Scatterpolar(r=df_p['Actual'], theta=df_p['Dimensión'], fill='toself', name='Hoy'))
            fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 10])), showlegend=True)
            st.plotly_chart(fig, use_container_width=True)
            st.table(df_p.style.format({"Actual": "{:.1f}", "Objetivo": "{:.1f}", "Brecha": "{:.1f}"}))

# --- TAB 3: REPORTE ---
with tab3:
    if not df_emp.empty:
        df_f_rep = interfaz_filtros(df_emp, "reporte")
        col_t, col_b = st.columns([4, 1])
        col_t.subheader("Resumen Ejecutivo")
        with col_b:
            w_bin = generar_word(emp_sel, df_f_rep, logo_file)
            st.download_button("📥 Word", data=w_bin, file_name=f"Estrategia_{emp_sel}.docx")
        
        # Mostrar tabla de compromisos incluyendo capacidades
        st.dataframe(df_f_rep[["Dimensión", "Foco", "Nombre", "Actual", "Brecha", "Capacidades_Distintivas", "Accion_1"]].sort_values(by="Brecha", ascending=False), use_container_width=True)